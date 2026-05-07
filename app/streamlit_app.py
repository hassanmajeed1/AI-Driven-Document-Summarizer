

import os, sys, time, re, tempfile, warnings
from pathlib import Path
from collections import defaultdict

import streamlit as st
import nltk, spacy
import plotly.graph_objects as go
import matplotlib, matplotlib.pyplot as plt
from wordcloud import WordCloud
from sklearn.feature_extraction.text import TfidfVectorizer
import numpy as np
import torch

warnings.filterwarnings("ignore")
matplotlib.use("Agg")

for _p in ["punkt", "stopwords", "punkt_tab"]:
    nltk.download(_p, quiet=True)

from nltk.tokenize import sent_tokenize, word_tokenize


CUDA_AVAILABLE = torch.cuda.is_available()
DEVICE_ID      = 0 if CUDA_AVAILABLE else -1

st.set_page_config(
    page_title="Document Intelligence System — Hassan Majeed",
    page_icon="◈",
    layout="wide",
    initial_sidebar_state="expanded",
)

C = {
    
    "bg":           "#050c14",   
    "surface":      "#0a1628",   
    "surface2":     "#0f1f38",   
    "surface3":     "#152844",   

  
    "border":       "#1a3050",   
    "border_glow":  "#00d4ff",   

    "accent":       "#00d4ff",   
    "accent2":      "#00aacc",   
    "accent3":      "#007fa3",
    "accent_dim":   "#00d4ff0f",   
    "accent_low":   "#00d4ff1a", 
    "accent_mid":   "#00d4ff2e",

   
    "text_bright":  "#e8f4fd",   
    "text_main":    "#7a9bb5",   
    "text_muted":   "#2e5070",   

    "green":        "#00e5b0",   
    "purple":       "#a78bfa",   
    "orange":       "#ff9f43",   
    "red":          "#ff6b6b",   
    "blue":         "#74b9ff",  

   
    "btn_text":     "#050c14",   
}


CSS = f"""
<style>
@import url('https://fonts.googleapis.com/css2?family=Inter:wght@300;400;500;600;700;800;900&family=JetBrains+Mono:wght@400;500&display=swap');

*, *::before, *::after {{ box-sizing: border-box; margin: 0; padding: 0; }}

html, body,
[data-testid="stAppViewContainer"],
[data-testid="stMain"],
.main {{
    background: {C["bg"]} !important;
    font-family: 'Inter', sans-serif !important;
}}

#MainMenu, footer, header,
[data-testid="stToolbar"],
[data-testid="stDecoration"] {{
    visibility: hidden !important;
    height: 0 !important;
}}

::-webkit-scrollbar {{ width: 4px; }}
::-webkit-scrollbar-track {{ background: {C["bg"]}; }}
::-webkit-scrollbar-thumb {{
    background: {C["border"]};
    border-radius: 2px;
    transition: background 0.3s;
}}
::-webkit-scrollbar-thumb:hover {{ background: {C["accent3"]}; }}


[data-testid="stSidebar"] {{
    background: {C["surface"]} !important;
    border-right: 1px solid {C["border"]} !important;
    padding-top: 0 !important;
}}

/* All sidebar text defaults */
[data-testid="stSidebar"],
[data-testid="stSidebar"] p,
[data-testid="stSidebar"] span,
[data-testid="stSidebar"] label,
[data-testid="stSidebar"] div {{
    color: {C["text_main"]} !important;
    font-family: 'Inter', sans-serif !important;
}}

/* Widget labels — clean muted */
[data-testid="stSidebar"] [data-testid="stWidgetLabel"] p,
[data-testid="stSidebar"] [data-testid="stWidgetLabel"] span {{
    color: {C["text_muted"]} !important;
    font-size: 0.72rem !important;
    font-weight: 500 !important;
    letter-spacing: 0.04em !important;
    text-transform: none !important;
}}


[data-testid="stSidebar"] [data-baseweb="slider"] [data-testid="stThumbValue"],
[data-testid="stSidebar"] [data-baseweb="slider"] [data-testid="stThumbValue"] *,
[data-testid="stSidebar"] [data-baseweb="slider"] [data-testid="stTickBarMin"],
[data-testid="stSidebar"] [data-baseweb="slider"] [data-testid="stTickBarMax"],
[data-testid="stSidebar"] [data-baseweb="slider"] [role="tooltip"],
[data-testid="stSidebar"] [data-baseweb="slider"] [aria-label*="thumb"] [data-testid="stThumbValue"] {{
    background: transparent !important;
    background-color: transparent !important;
    border: none !important;
    box-shadow: none !important;
    color: #00d4ff !important;
    font-weight: 700 !important;
    font-size: 0.78rem !important;
    padding: 0 !important;
    margin: 0 !important;
    width: auto !important;
    height: auto !important;
    transform: none !important;
    backdrop-filter: none !important;
}}

/* Slider min/max range labels */
[data-testid="stSidebar"] [data-testid="stSlider"] [data-testid="stTickBarMin"],
[data-testid="stSidebar"] [data-testid="stSlider"] [data-testid="stTickBarMax"],
[data-testid="stSidebar"] [data-testid="stSlider"] span:not([data-testid="stThumbValue"]) {{
    color: {C["text_muted"]} !important;
    font-size: 0.72rem !important;
}}

/* Slider track */
[data-testid="stSidebar"] [role="slider"] {{
    background: {C["accent"]} !important;
    box-shadow: 0 0 8px {C["accent"]}66 !important;
}}
[data-testid="stSidebar"] [data-baseweb="slider"] div[data-testid] {{
    background: {C["border"]} !important;
}}

/* Selectbox */
[data-testid="stSidebar"] [data-baseweb="select"] > div {{
    background: {C["surface2"]} !important;
    border: 1px solid {C["border"]} !important;
    border-radius: 8px !important;
    color: {C["text_bright"]} !important;
    transition: border-color 0.25s !important;
}}
[data-testid="stSidebar"] [data-baseweb="select"] > div:hover {{
    border-color: {C["accent"]}55 !important;
}}
[data-testid="stSidebar"] [data-baseweb="select"] [data-testid="stMarkdownContainer"] p {{
    color: {C["text_bright"]} !important;
    font-size: 0.85rem !important;
}}

/* Radio buttons */
[data-testid="stSidebar"] [data-testid="stRadio"] > div {{
    gap: 6px !important;
    display: flex !important;
    flex-direction: column !important;
}}
[data-testid="stSidebar"] [data-testid="stRadio"] label {{
    background: {C["surface2"]} !important;
    border: 1px solid {C["border"]} !important;
    border-radius: 8px !important;
    padding: 9px 12px !important;
    transition: all 0.2s ease !important;
    cursor: pointer !important;
    width: 100% !important;
}}
[data-testid="stSidebar"] [data-testid="stRadio"] label:hover {{
    border-color: {C["accent"]}44 !important;
    background: {C["surface3"]} !important;
}}
[data-testid="stSidebar"] [data-testid="stRadio"] label:has(input:checked) {{
    border-color: {C["accent"]}66 !important;
    background: {C["accent_low"]} !important;
    box-shadow: 0 0 0 1px {C["accent"]}22 inset !important;
}}
[data-testid="stSidebar"] [data-testid="stRadio"] label p,
[data-testid="stSidebar"] [data-testid="stRadio"] label span {{
    color: {C["text_main"]} !important;
    font-size: 0.84rem !important;
    letter-spacing: 0 !important;
    text-transform: none !important;
}}
[data-testid="stSidebar"] [data-testid="stRadio"] label:has(input:checked) p {{
    color: {C["text_bright"]} !important;
}}


.hero-wrap {{
    background: linear-gradient(135deg, {C["surface"]} 0%, {C["surface2"]} 50%, {C["bg"]} 100%);
    border: 1px solid {C["border"]};
    border-radius: 20px;
    padding: 52px 56px 44px;
    margin-bottom: 28px;
    position: relative;
    overflow: hidden;
}}
.hero-wrap::before {{
    content: '';
    position: absolute;
    top: -120px; right: -80px;
    width: 400px; height: 400px;
    border-radius: 50%;
    background: radial-gradient(circle, {C["accent"]}0f 0%, transparent 60%);
    pointer-events: none;
}}
.hero-wrap::after {{
    content: '';
    position: absolute;
    bottom: -60px; left: 10%;
    width: 300px; height: 300px;
    border-radius: 50%;
    background: radial-gradient(circle, {C["accent2"]}08 0%, transparent 65%);
    pointer-events: none;
}}
.hero-eyebrow {{
    display: flex; align-items: center; gap: 12px;
    margin-bottom: 20px;
}}
.hero-eyebrow-line {{
    width: 32px; height: 1.5px;
    background: linear-gradient(90deg, {C["accent"]}, transparent);
}}
.hero-eyebrow-text {{
    font-size: 0.65rem; font-weight: 600;
    letter-spacing: 0.22em; text-transform: uppercase;
    color: {C["accent"]};
}}
.hero-title {{
    font-family: 'Inter', sans-serif;
    font-size: clamp(1.8rem, 3.2vw, 3rem);
    font-weight: 900;
    letter-spacing: -0.03em;
    line-height: 1.06;
    color: {C["text_bright"]};
    margin-bottom: 14px;
}}
.hero-title .a {{
    background: linear-gradient(90deg, {C["accent"]}, {C["blue"]});
    -webkit-background-clip: text;
    -webkit-text-fill-color: transparent;
    background-clip: text;
}}
.hero-sub {{
    font-size: 0.92rem; font-weight: 400;
    color: {C["text_main"]}; line-height: 1.75;
    max-width: 540px; margin-bottom: 28px;
}}
.hero-byline {{ display: flex; align-items: center; gap: 14px; }}
.hero-dot {{
    width: 7px; height: 7px; border-radius: 50%;
    background: {C["accent"]};
    box-shadow: 0 0 10px {C["accent"]};
}}
.hero-byline-text {{
    font-size: 0.74rem; font-weight: 500;
    letter-spacing: 0.06em; color: {C["text_muted"]};
}}


.sec-label {{
    display: flex; align-items: center; gap: 10px;
    margin-bottom: 8px;
}}
.sec-line {{
    width: 3px; height: 18px;
    background: linear-gradient(180deg, {C["accent"]}, {C["accent2"]});
    border-radius: 2px;
}}
.sec-eyebrow {{
    font-size: 0.62rem; font-weight: 700;
    letter-spacing: 0.18em; text-transform: uppercase;
    color: {C["accent"]};
}}
.sec-title {{
    font-family: 'Inter', sans-serif;
    font-size: 1.2rem; font-weight: 800;
    letter-spacing: -0.025em;
    color: {C["text_bright"]};
    margin-bottom: 20px;
}}


.m-card {{
    background: {C["surface"]};
    border: 1px solid {C["border"]};
    border-radius: 16px;
    padding: 22px 16px 18px;
    position: relative;
    overflow: hidden;
    height: 100%;
    transition: border-color 0.3s ease, transform 0.3s ease, box-shadow 0.3s ease;
    cursor: default;
}}
.m-card::before {{
    content: '';
    position: absolute; top: 0; left: 0; right: 0;
    height: 2px;
    background: linear-gradient(90deg, {C["accent"]}, {C["accent2"]}, transparent);
    opacity: 0;
    transition: opacity 0.3s ease;
}}
.m-card:hover {{
    border-color: {C["accent"]}44;
    transform: translateY(-4px);
    box-shadow: 0 12px 32px {C["accent"]}14, 0 4px 12px {C["bg"]}80;
}}
.m-card:hover::before {{ opacity: 1; }}
.m-num {{
    font-family: 'Inter', sans-serif;
    font-size: clamp(0.95rem, 2vw, 1.7rem);
    font-weight: 800;
    color: {C["accent"]};
    line-height: 1.1;
    margin-bottom: 7px;
    word-break: break-all;
    overflow-wrap: break-word;
    white-space: nowrap;
    overflow: hidden;
    text-overflow: ellipsis;
}}
.m-lbl {{
    font-size: 0.6rem; font-weight: 600;
    color: {C["text_muted"]};
    text-transform: uppercase;
    letter-spacing: 0.1em;
    line-height: 1.3;
}}


.sum-card {{
    background: {C["surface"]};
    border: 1px solid {C["border"]};
    border-radius: 16px;
    padding: 28px 32px;
    margin-bottom: 14px;
    position: relative;
    overflow: hidden;
    transition: border-color 0.3s, box-shadow 0.3s;
}}
.sum-card::before {{
    content: '';
    position: absolute; top: 0; left: 0;
    width: 3px; height: 100%;
    background: linear-gradient(180deg, {C["accent"]}, {C["accent2"]}, transparent);
    border-radius: 0 2px 2px 0;
}}
.sum-card:hover {{
    border-color: {C["accent"]}33;
    box-shadow: 0 8px 30px {C["accent"]}0a;
}}
.sum-badge {{
    display: inline-block;
    font-size: 0.6rem; font-weight: 700;
    letter-spacing: 0.12em; text-transform: uppercase;
    padding: 3px 10px;
    background: {C["accent_dim"]};
    color: {C["accent"]};
    border-radius: 20px;
    border: 1px solid {C["accent"]}2e;
    margin-bottom: 10px;
}}
.sum-head {{
    font-family: 'Inter', sans-serif;
    font-size: 0.88rem; font-weight: 700;
    color: {C["text_bright"]}; margin-bottom: 12px;
}}
.sum-text {{
    font-size: 0.91rem; font-weight: 400;
    color: {C["text_main"]}; line-height: 1.82;
}}


.chips-wrap {{ display: flex; flex-wrap: wrap; gap: 8px; margin-top: 6px; }}
.chip-h {{
    display: inline-block;
    font-size: 0.72rem; font-weight: 600;
    padding: 5px 14px;
    background: {C["accent_dim"]};
    color: {C["accent"]};
    border: 1px solid {C["accent"]}33;
    border-radius: 20px;
    letter-spacing: 0.02em;
    transition: background 0.2s, border-color 0.2s, box-shadow 0.2s;
    cursor: default;
}}
.chip-h:hover {{
    background: {C["accent_low"]};
    border-color: {C["accent"]}55;
    box-shadow: 0 0 10px {C["accent"]}22;
}}
.chip-m {{
    display: inline-block;
    font-size: 0.72rem; font-weight: 400;
    padding: 5px 14px;
    background: {C["surface2"]};
    color: {C["text_main"]};
    border: 1px solid {C["border"]};
    border-radius: 20px;
    transition: border-color 0.2s;
    cursor: default;
}}
.chip-m:hover {{ border-color: {C["text_muted"]}; }}
.chip-label {{
    font-size: 0.6rem; font-weight: 700;
    letter-spacing: 0.14em; text-transform: uppercase;
    color: {C["text_muted"]}; margin-bottom: 8px; margin-top: 14px;
}}


.ent-sec {{ margin-bottom: 22px; }}
.ent-type {{
    font-size: 0.6rem; font-weight: 700;
    letter-spacing: 0.16em; text-transform: uppercase;
    margin-bottom: 9px;
}}
.ent-pills {{ display: flex; flex-wrap: wrap; gap: 6px; }}
.ent-pill {{
    display: inline-block;
    font-size: 0.75rem; font-weight: 400;
    padding: 4px 12px;
    background: {C["surface2"]};
    color: {C["text_bright"]};
    border: 1px solid {C["border"]};
    border-radius: 8px;
    transition: border-color 0.2s, background 0.2s;
    cursor: default;
}}
.ent-pill:hover {{
    background: {C["surface3"]};
    border-color: {C["border_glow"]}33;
}}


.slim-div {{
    height: 1px;
    background: linear-gradient(90deg, transparent, {C["border"]}, transparent);
    margin: 28px 0;
}}


.stProgress > div > div > div > div {{
    background: linear-gradient(90deg, {C["accent3"]}, {C["accent"]}) !important;
    border-radius: 4px !important;
    box-shadow: 0 0 12px {C["accent"]}44 !important;
}}
.stProgress > div > div {{
    background: {C["surface2"]} !important;
    border-radius: 4px !important;
    border: 1px solid {C["border"]} !important;
}}

.stButton > button {{
    background: linear-gradient(135deg, {C["accent"]} 0%, {C["accent2"]} 100%) !important;
    color: {C["btn_text"]} !important;       /* ← FIX #4: dark text */
    border: none !important;
    border-radius: 10px !important;
    font-family: 'Inter', sans-serif !important;
    font-size: 0.8rem !important;
    font-weight: 800 !important;
    letter-spacing: 0.1em !important;
    text-transform: uppercase !important;
    padding: 14px 28px !important;
    box-shadow: 0 4px 24px {C["accent"]}44 !important;
    transition: all 0.25s cubic-bezier(0.23, 1, 0.32, 1) !important;
    width: 100% !important;
}}
.stButton > button * {{
    color: {C["btn_text"]} !important;       
    font-weight: 800 !important;
}}
.stButton > button:hover {{
    transform: translateY(-2px) !important;
    box-shadow: 0 12px 40px {C["accent"]}55, 0 0 0 1px {C["accent"]}44 !important;
    filter: brightness(1.06) !important;
}}
.stButton > button:active {{ transform: translateY(0) !important; }}


[data-testid="stDownloadButton"] > button {{
    background: {C["surface2"]} !important;
    color: {C["accent"]} !important;
    border: 1px solid {C["border"]} !important;
    border-radius: 10px !important;
    font-family: 'Inter', sans-serif !important;
    font-size: 0.75rem !important;
    font-weight: 700 !important;
    letter-spacing: 0.08em !important;
    text-transform: uppercase !important;
    padding: 11px 16px !important;
    width: 100% !important;
    transition: all 0.25s ease !important;
    box-shadow: none !important;
}}
[data-testid="stDownloadButton"] > button:hover {{
    background: {C["accent_low"]} !important;
    border-color: {C["accent"]}55 !important;
    transform: translateY(-2px) !important;
    box-shadow: 0 6px 20px {C["accent"]}1a !important;
}}


[data-testid="stFileUploader"] {{
    background: {C["surface"]} !important;
    border: 1.5px dashed {C["border"]} !important;
    border-radius: 14px !important;
    padding: 10px !important;
    transition: border-color 0.3s, box-shadow 0.3s !important;
}}
[data-testid="stFileUploader"]:hover {{
    border-color: {C["accent"]}44 !important;
    box-shadow: 0 0 0 3px {C["accent"]}08 !important;
}}
[data-testid="stFileUploader"] * {{ color: {C["text_main"]} !important; }}


.stTextArea textarea {{
    background: {C["surface"]} !important;
    border: 1px solid {C["border"]} !important;
    border-radius: 12px !important;
    color: {C["text_bright"]} !important;
    font-family: 'Inter', sans-serif !important;
    font-size: 0.9rem !important;
    line-height: 1.7 !important;
    transition: border-color 0.25s, box-shadow 0.25s !important;
}}
.stTextArea textarea:focus {{
    border-color: {C["accent"]}66 !important;
    box-shadow: 0 0 0 3px {C["accent"]}12 !important;
    outline: none !important;
}}
.stTextArea label {{ color: {C["text_muted"]} !important; }}


[data-baseweb="tab-list"] {{
    background: {C["surface"]} !important;
    border-radius: 12px !important;
    padding: 4px !important;
    gap: 4px !important;
    border: 1px solid {C["border"]} !important;
}}
[data-baseweb="tab"] {{
    font-family: 'Inter', sans-serif !important;
    font-size: 0.74rem !important;
    font-weight: 700 !important;
    letter-spacing: 0.08em !important;
    text-transform: uppercase !important;
    color: {C["text_muted"]} !important;
    border-radius: 9px !important;
    padding: 9px 18px !important;
    border: none !important;
    transition: color 0.2s, background 0.2s !important;
}}
[data-baseweb="tab"]:hover:not([aria-selected="true"]) {{
    color: {C["text_main"]} !important;
    background: {C["surface2"]} !important;
}}
/* Active tab — bright background, DARK text (Fix #6) */
[aria-selected="true"][data-baseweb="tab"] {{
    background: linear-gradient(135deg, {C["accent"]}, {C["accent2"]}) !important;
    color: {C["btn_text"]} !important;       /* ← FIX #6 */
    box-shadow: 0 2px 12px {C["accent"]}44 !important;
}}
[aria-selected="true"][data-baseweb="tab"] *,
[aria-selected="true"][data-baseweb="tab"] p,
[aria-selected="true"][data-baseweb="tab"] span {{
    color: {C["btn_text"]} !important;       /* ← FIX #6: force all children */
}}
[data-testid="stTabContent"] {{ padding-top: 16px !important; }}


.stAlert {{ border-radius: 12px !important; border: none !important; }}
[data-testid="stSuccessMessage"] {{
    background: {C["green"]}0d !important;
    border: 1px solid {C["green"]}2a !important;
    border-left: 3px solid {C["green"]} !important;
    border-radius: 0 12px 12px 0 !important;
    color: {C["green"]} !important;
}}
[data-testid="stInfoMessage"] {{
    background: {C["accent_dim"]} !important;
    border-left: 3px solid {C["accent"]} !important;
    color: {C["accent"]} !important;
}}
[data-testid="stWarningMessage"] {{
    background: {C["orange"]}0d !important;
    border-left: 3px solid {C["orange"]} !important;
    color: {C["orange"]} !important;
}}
[data-testid="stErrorMessage"] {{
    background: {C["red"]}0d !important;
    border-left: 3px solid {C["red"]} !important;
    color: {C["red"]} !important;
}}

[data-testid="stExpander"] {{
    background: {C["surface"]} !important;
    border: 1px solid {C["border"]} !important;
    border-radius: 12px !important;
    transition: border-color 0.3s !important;
}}
[data-testid="stExpander"]:hover {{
    border-color: {C["accent"]}33 !important;
}}


.stSpinner > div {{
    border-color: {C["accent"]} transparent transparent transparent !important;
}}


.main h1, .main h2, .main h3 {{
    font-family: 'Inter', sans-serif !important;
    color: {C["text_bright"]} !important;
    font-weight: 800 !important;
}}
.main p {{ color: {C["text_main"]} !important; line-height: 1.7 !important; }}
.main code {{
    background: {C["surface2"]} !important;
    color: {C["accent"]} !important;
    border-radius: 5px !important;
    padding: 2px 7px !important;
    font-family: 'JetBrains Mono', monospace !important;
    font-size: 0.83em !important;
    border: 1px solid {C["border"]} !important;
}}
</style>
"""

st.markdown(CSS, unsafe_allow_html=True)


@st.cache_resource(show_spinner=False)
def load_spacy():
    try:
        return spacy.load("en_core_web_sm")
    except OSError:
        import subprocess
        subprocess.run(
            [sys.executable, "-m", "spacy", "download", "en_core_web_sm"],
            check=True,
        )
        return spacy.load("en_core_web_sm")


@st.cache_resource(show_spinner=False)
def load_keybert():
    from keybert import KeyBERT
    return KeyBERT(model="all-MiniLM-L6-v2")


@st.cache_resource(show_spinner=False)
def load_summarizer(model_key: str):
    
    from transformers import pipeline, AutoTokenizer, AutoModelForSeq2SeqLM

    names = {
        "fast":    "sshleifer/distilbart-cnn-12-6",
        "quality": "facebook/bart-large-cnn",
    }
    model_name = names.get(model_key, names["fast"])
    tokenizer  = AutoTokenizer.from_pretrained(model_name)

    if CUDA_AVAILABLE:
        model = AutoModelForSeq2SeqLM.from_pretrained(
            model_name, torch_dtype=torch.float16
        ).to("cuda")
        pipe = pipeline(
            "summarization",
            model=model,
            tokenizer=tokenizer,
            device=0,
            torch_dtype=torch.float16,
        )
    else:
        pipe = pipeline(
            "summarization",
            model=model_name,
            tokenizer=tokenizer,
            device=-1,
        )
    return pipe, tokenizer




def _extract_pdf(path: str) -> str:
    import pdfplumber, PyPDF2
    parts = []
    try:
        with pdfplumber.open(path) as pdf:
            for pg in pdf.pages:
                t = pg.extract_text()
                if t:
                    parts.append(t)
        if parts:
            return "\n\n".join(parts)
    except Exception:
        pass
    parts = []
    with open(path, "rb") as f:
        for pg in PyPDF2.PdfReader(f).pages:
            parts.append(pg.extract_text() or "")
    return "\n\n".join(parts)


def _extract_docx(path: str) -> str:
    from docx import Document
    doc   = Document(path)
    parts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    for tbl in doc.tables:
        for row in tbl.rows:
            r = " | ".join(c.text.strip() for c in row.cells if c.text.strip())
            if r:
                parts.append(r)
    return "\n".join(parts)


def _extract_txt(path: str) -> str:
    for enc in ["utf-8", "latin-1", "cp1252"]:
        try:
            return open(path, "r", encoding=enc).read()
        except UnicodeDecodeError:
            continue
    raise ValueError("Cannot decode file.")


def extract_text(uploaded_file) -> str:
    suffix = Path(uploaded_file.name).suffix.lower()
    with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
        tmp.write(uploaded_file.getvalue())
        path = tmp.name
    try:
        fn = {
            ".pdf":  _extract_pdf,
            ".docx": _extract_docx,
            ".doc":  _extract_docx,
            ".txt":  _extract_txt,
        }
        if suffix not in fn:
            raise ValueError(f"Unsupported format: {suffix}")
        return fn[suffix](path)
    finally:
        os.unlink(path)




def clean_text(text: str) -> str:
    text = re.sub(r"\n{3,}", "\n\n", text)
    text = re.sub(r" {2,}", " ", text)
    for o, n in {
        "\u2019": "'", "\u2018": "'", "\u201c": '"', "\u201d": '"',
        "\u2013": "-", "\u2014": "-", "\xa0": " ",
    }.items():
        text = text.replace(o, n)
    text = re.sub(r"Page \d+ of \d+", "", text, flags=re.IGNORECASE)
    text = re.sub(r"[-_=]{4,}", "", text)
    return "\n".join(l.strip() for l in text.split("\n") if l.strip())


def get_sentences(text: str) -> list:
    return [s.strip() for s in sent_tokenize(text) if len(s.strip()) > 20]


def get_stats(text: str) -> dict:
    words = [w for w in word_tokenize(text.lower()) if w.isalpha()]
    sents = get_sentences(text)
    return {
        "words":      len(words),
        "sentences":  len(sents),
        "unique":     len(set(words)),
        "characters": len(text),
        "read_min":   round(len(words) / 200, 1),
    }




def _safe_chunks(tokenizer, text: str, max_tokens: int = 900) -> list:
   
    ids    = tokenizer.encode(text, add_special_tokens=False)
    chunks = []
    for i in range(0, len(ids), max_tokens):
        chunk_ids = ids[i: i + max_tokens]
        chunks.append(tokenizer.decode(chunk_ids, skip_special_tokens=True))
    return chunks if chunks else [text]


def extractive_summary(text: str, sentences: list, n: int = 5) -> str:
    if len(sentences) <= n:
        return " ".join(sentences)
    try:
        vec = TfidfVectorizer(stop_words="english", ngram_range=(1, 2))
        mat = vec.fit_transform(sentences)
        scores  = np.array(mat.sum(axis=1)).flatten()
        top_idx = sorted(np.argsort(scores)[-n:].tolist())
        return " ".join(sentences[i] for i in top_idx)
    except Exception:
        return " ".join(sentences[:n])


def abstractive_summary(pipe, tokenizer, text: str,
                        max_len: int = 150, min_len: int = 50) -> str:
    min_len = min(min_len, max_len - 10)
    chunks  = _safe_chunks(tokenizer, text, max_tokens=900)

    summaries = []
    for chunk in chunks:
        if not chunk.strip():
            continue
        try:
            out = pipe(
                chunk,
                max_length=max_len,
                min_length=min_len,
                do_sample=False,
                truncation=True,
            )
            summaries.append(out[0]["summary_text"])
        except Exception as e:
            summaries.append(f"[Chunk error: {e}]")

    if not summaries:
        return "Summary could not be generated."
    if len(summaries) == 1:
        return summaries[0]

    combined = " ".join(summaries)
    try:
        final = pipe(
            combined,
            max_length=max_len,
            min_length=min_len,
            do_sample=False,
            truncation=True,
        )
        return final[0]["summary_text"]
    except Exception:
        return combined




def extract_keywords(kw_model, text: str, top_n: int = 12) -> list:
    return kw_model.extract_keywords(
        text,
        keyphrase_ngram_range=(1, 3),
        stop_words="english",
        use_mmr=True,
        diversity=0.5,
        top_n=top_n,
    )




ENTITY_META = {
    "ORG":     ("Organizations",  C["accent"]),
    "PERSON":  ("People",         C["blue"]),
    "MONEY":   ("Financial",      C["green"]),
    "DATE":    ("Dates",          C["purple"]),
    "GPE":     ("Locations",      C["orange"]),
    "PERCENT": ("Percentages",    C["red"]),
    "PRODUCT": ("Products",       C["text_main"]),
}


def extract_entities(nlp_model, text: str) -> dict:
    doc    = nlp_model(text[:900_000])
    result = defaultdict(set)
    for ent in doc.ents:
        if ent.label_ in ENTITY_META and len(ent.text.strip()) > 1:
            result[ent.label_].add(ent.text.strip())
    return {k: sorted(v) for k, v in result.items()}




def keyword_bar_chart(keywords):
    words  = [kw[0] for kw in keywords[:12]]
    scores = [kw[1] for kw in keywords[:12]]
    colorscale = [
        [0.0, C["surface3"]],
        [0.5, C["accent3"]],
        [1.0, C["accent"]],
    ]
    fig = go.Figure(go.Bar(
        x=scores, y=words, orientation="h",
        marker=dict(
            color=scores,
            colorscale=colorscale,
            showscale=False,
            line=dict(width=0),
        ),
        hovertemplate="<b>%{y}</b><br>Score: %{x:.3f}<extra></extra>",
    ))
    fig.update_layout(
        plot_bgcolor=C["surface"],
        paper_bgcolor=C["surface"],
        font=dict(family="Inter", color=C["text_main"], size=11),
        margin=dict(l=10, r=20, t=20, b=40),
        height=370,
        xaxis=dict(
            showgrid=True,
            gridcolor=C["border"],
            zeroline=False,
            tickfont=dict(color=C["text_muted"]),
            title=dict(text="Relevance Score",
                       font=dict(color=C["text_muted"], size=10)),
        ),
        yaxis=dict(
            showgrid=False,
            tickfont=dict(color=C["text_bright"], size=11),
        ),
    )
    return fig


def wordcloud_fig(keywords):
    freq = {w: s for w, s in keywords}
    wc   = WordCloud(
        width=720, height=270,
        background_color=C["surface"],
        colormap="cool",
        max_words=30,
        prefer_horizontal=0.8,
    ).generate_from_frequencies(freq)
    fig, ax = plt.subplots(figsize=(9, 3))
    fig.patch.set_facecolor(C["surface"])
    ax.set_facecolor(C["surface"])
    ax.imshow(wc, interpolation="bilinear")
    ax.axis("off")
    plt.tight_layout(pad=0)
    return fig




def build_report(stats, ab_sum, ex_sum, keywords, entities) -> str:
    kw_block  = "\n".join(f"  - {k}  (score: {v:.3f})" for k, v in keywords[:15])
    ent_block = ""
    for label, items in entities.items():
        ent_block += f"\n### {ENTITY_META.get(label, (label,))[0]}\n"
        ent_block += "\n".join(f"  - {i}" for i in items[:8])
    return f"""# Document Intelligence Report
**System:** AI-Powered Document Intelligence System using Transformer Models
**Author:** Hassan Majeed

---

## Document Statistics
| Metric | Value |
|--------|-------|
| Words | {stats['words']} |
| Sentences | {stats['sentences']} |
| Unique Words | {stats['unique']} |
| Characters | {stats['characters']} |
| Reading Time | {stats['read_min']} min |

---

## Transformer Summary (Abstractive)
{ab_sum or '_Not generated._'}

---

## Key Sentences (Extractive)
{ex_sum or '_Not generated._'}

---

## Top Keywords
{kw_block or '_None._'}

---

## Named Entities
{ent_block or '_None detected._'}

---
*Generated with HuggingFace Transformers · KeyBERT · SpaCy*
"""



def render_hero():
    st.markdown(f"""
    <div class="hero-wrap">
        <div class="hero-eyebrow">
            <div class="hero-eyebrow-line"></div>
            <span class="hero-eyebrow-text">
                NLP &nbsp;·&nbsp; Transformer Models &nbsp;·&nbsp; Document Intelligence
            </span>
        </div>
        <div class="hero-title">
            AI-Powered Document<br>
            <span class="a">Intelligence System</span>
        </div>
        <div class="hero-sub">
            Upload any business document and extract structured insights —
            transformer summaries, semantic keywords, and named entities,
            all delivered in a single unified pipeline.
        </div>
        <div class="hero-byline">
            <div class="hero-dot"></div>
            <span class="hero-byline-text">
                Built by Hassan Majeed &nbsp;·&nbsp; using Transformer Models
            </span>
        </div>
    </div>
    """, unsafe_allow_html=True)


def sec_label(eyebrow: str, title: str):
    st.markdown(f"""
    <div class="sec-label">
        <div class="sec-line"></div>
        <span class="sec-eyebrow">{eyebrow}</span>
    </div>
    <div class="sec-title">{title}</div>
    """, unsafe_allow_html=True)


def divider():
    st.markdown('<div class="slim-div"></div>', unsafe_allow_html=True)


def metric_cards(stats: dict):
    items = [
        (str(stats["words"]),        "Words"),
        (str(stats["sentences"]),    "Sentences"),
        (str(stats["unique"]),       "Unique Words"),
        (f"{stats['read_min']} min", "Read Time"),
        (str(stats["characters"]),   "Characters"),
    ]
    cols = st.columns(5, gap="small")
    for col, (val, lbl) in zip(cols, items):
        with col:
            st.markdown(
                f'<div class="m-card">'
                f'<div class="m-num">{val}</div>'
                f'<div class="m-lbl">{lbl}</div>'
                f'</div>',
                unsafe_allow_html=True,
            )
    st.markdown("<div style='margin-bottom:8px'></div>", unsafe_allow_html=True)


def summary_block(title: str, badge: str, text: str):
    st.markdown(f"""
    <div class="sum-card">
        <div class="sum-badge">{badge}</div>
        <div class="sum-head">{title}</div>
        <div class="sum-text">{text}</div>
    </div>
    """, unsafe_allow_html=True)


def keyword_chips(keywords: list):
    high = [(k, s) for k, s in keywords if s >= 0.55]
    rest = [(k, s) for k, s in keywords if s <  0.55]
    if high:
        st.markdown('<div class="chip-label">High Relevance</div>',
                    unsafe_allow_html=True)
        chips = "".join(f'<span class="chip-h">{k}</span>' for k, _ in high)
        st.markdown(f'<div class="chips-wrap">{chips}</div>', unsafe_allow_html=True)
    if rest:
        st.markdown('<div class="chip-label">Supporting Terms</div>',
                    unsafe_allow_html=True)
        chips = "".join(f'<span class="chip-m">{k}</span>' for k, _ in rest)
        st.markdown(f'<div class="chips-wrap">{chips}</div>', unsafe_allow_html=True)


def entity_display(entities: dict):
    if not entities:
        st.info("No significant named entities detected.")
        return
    keys   = list(entities.keys())
    col_l, col_r = st.columns(2, gap="large")
    groups = [keys[:4], keys[4:]]
    for col, grp in zip([col_l, col_r], groups):
        with col:
            for label in grp:
                items      = entities[label]
                lbl_name, color = ENTITY_META.get(label, (label, C["text_main"]))
                pills = "".join(
                    f'<span class="ent-pill">{i}</span>' for i in items[:8]
                )
                st.markdown(
                    f'<div class="ent-sec">'
                    f'<div class="ent-type" style="color:{color}">{lbl_name}</div>'
                    f'<div class="ent-pills">{pills}</div>'
                    f'</div>',
                    unsafe_allow_html=True,
                )




def render_sidebar() -> dict:
    with st.sidebar:
        # Header
        st.markdown(f"""
        <div style="padding: 28px 0 4px;">
            <div style="font-family:'Inter',sans-serif; font-size:1rem;
                        font-weight:900; color:{C['text_bright']};
                        letter-spacing:-0.02em; margin-bottom:6px">
                Settings
            </div>
            <div style="width:24px; height:2px;
                        background:linear-gradient(90deg,{C['accent']},{C['accent2']});
                        border-radius:1px; margin-bottom:24px">
            </div>
        </div>
        """, unsafe_allow_html=True)

        def sb_head(txt: str):
            """Sidebar section heading — muted, no yellow."""
            st.markdown(
                f"<div style='font-size:.6rem; font-weight:700; "
                f"letter-spacing:.16em; text-transform:uppercase; "
                f"color:{C['text_muted']}; margin:18px 0 8px'>{txt}</div>",
                unsafe_allow_html=True,
            )

        sb_head("Input Method")
        input_method = st.radio(
            "input_method",
            ["Upload File", "Paste Text"],
            label_visibility="collapsed",
        )

        sb_head("Summary")
        summary_type = st.selectbox(
            "Summary Type",
            ["Both (Recommended)", "Transformer Only", "Extractive Only"],
        )
        n_sentences = st.slider("Sentence count", 2, 10, 5)
        summary_length = st.select_slider(
            "Output length",
            options=["Short", "Medium", "Long"],
            value="Medium",
        )
        length_map = {"Short": (30, 80), "Medium": (60, 150), "Long": (100, 260)}
        min_len, max_len = length_map[summary_length]

        sb_head("Keywords")
        n_keywords = st.slider("Keyword count", 5, 25, 12)

        sb_head("Model")
        model_choice = st.radio(
            "model_choice",
            ["Fast  —  DistilBART ", "Quality  —  BART-large "],
            label_visibility="collapsed",
        )
        model_key = "fast" if "Fast" in model_choice else "quality"

     
        st.markdown(
            f"<div style='margin-top:20px; padding:14px 16px; "
            f"background:{C['surface2']}; border:1px solid {C['border']}; "
            f"border-radius:10px; font-size:.73rem; "
            f"color:{C['text_muted']}; line-height:1.65'>"
            f"Models are cached after first load.<br>"
            f"First run downloads weights from HuggingFace."
            f"</div>",
            unsafe_allow_html=True,
        )

    return dict(
        input_method=input_method,
        summary_type=summary_type,
        n_sentences=n_sentences,
        min_len=min_len,
        max_len=max_len,
        n_keywords=n_keywords,
        model_key=model_key,
    )


def run_analysis(raw_text: str, cfg: dict):
    bar = st.progress(0, text="Starting…")
    t0  = time.time()

    bar.progress(8, "Cleaning and tokenising…")
    clean = clean_text(raw_text)
    sents = get_sentences(clean)
    stats = get_stats(clean)

    bar.progress(16, "Loading NLP models (cached after first run)…")
    nlp_model        = load_spacy()
    kw_model         = load_keybert()
    pipe = tokenizer = None
    if cfg["summary_type"] in ("Both (Recommended)", "Transformer Only"):
        pipe, tokenizer = load_summarizer(cfg["model_key"])

    bar.progress(36, "Extracting keywords with KeyBERT…")
    keywords = extract_keywords(kw_model, clean, top_n=cfg["n_keywords"])

    bar.progress(52, "Running Named Entity Recognition…")
    entities = extract_entities(nlp_model, clean)

    ex_sum = ab_sum = None
    if cfg["summary_type"] in ("Both (Recommended)", "Extractive Only"):
        bar.progress(64, "Computing TF-IDF extractive summary…")
        ex_sum = extractive_summary(clean, sents, cfg["n_sentences"])

    if pipe is not None:
        bar.progress(76, "Generating transformer summary…")
        ab_sum = abstractive_summary(
            pipe, tokenizer, clean,
            max_len=cfg["max_len"],
            min_len=cfg["min_len"],
        )

    bar.progress(100, "Done.")
    bar.empty()

    divider()

    sec_label("Document Statistics", "Document at a glance")
    metric_cards(stats)
    divider()

    sec_label("Summaries", "Document summaries")
    if ab_sum and ex_sum:
        t1, t2 = st.tabs(["Transformer Summary", "Key Sentences"])
        with t1:
            summary_block("Transformer-Generated Summary", "BART Abstractive", ab_sum)
        with t2:
            summary_block("Most Important Sentences", "TF-IDF Extractive", ex_sum)
    elif ab_sum:
        summary_block("Transformer-Generated Summary", "BART Abstractive", ab_sum)
    elif ex_sum:
        summary_block("Most Important Sentences", "TF-IDF Extractive", ex_sum)
    divider()

    sec_label("Keyword Extraction", "Key topics & phrases")
    col_chips, col_chart = st.columns([1, 1], gap="large")
    with col_chips:
        keyword_chips(keywords)
    with col_chart:
        st.plotly_chart(
            keyword_bar_chart(keywords),
            use_container_width=True,
            config={"displayModeBar": False},
        )
    st.markdown(
        f"<div style='font-size:.6rem; font-weight:700; letter-spacing:.16em; "
        f"text-transform:uppercase; color:{C['accent']}; "
        f"margin:8px 0 10px'>Keyword Cloud</div>",
        unsafe_allow_html=True,
    )
    st.pyplot(wordcloud_fig(keywords), use_container_width=True)
    divider()

    sec_label("Named Entity Recognition", "Entities detected")
    entity_display(entities)
    divider()

    
    sec_label("Export", "Download analysis report")
    report_md = build_report(stats, ab_sum, ex_sum, keywords, entities)

    rtf = (
        r"{\rtf1\ansi\deff0{\fonttbl{\f0 Arial;}}\f0\fs22 "
        + report_md
            .replace("\\", "\\\\")
            .replace("\n",  r"\par ")
            .replace("{",   r"\{")
            .replace("}",   r"\}")
        + r"}"
    )

    c1, c2, c3 = st.columns(3, gap="medium")
    with c1:
        st.download_button(
            "Download Markdown",
            data=report_md,
            file_name="report_hassan_majeed.md",
            mime="text/markdown",
        )
    with c2:
        st.download_button(
            "Download Word (.doc)",
            data=rtf.encode("utf-8"),
            file_name="report_hassan_majeed.doc",
            mime="application/msword",
        )
    with c3:
        st.download_button(
            "Download Clean Text",
            data=clean.encode("utf-8"),
            file_name="cleaned_document.txt",
            mime="text/plain",
        )




def main():
    render_hero()
    cfg = render_sidebar()

    raw_text = None

    if cfg["input_method"] == "Upload File":
        uploaded = st.file_uploader(
            "Drop your document — PDF, Word, or plain text",
            type=["pdf", "docx", "txt"],
            label_visibility="collapsed",
        )
        if uploaded:
            with st.spinner("Reading document…"):
                try:
                    raw_text = extract_text(uploaded)
                except Exception as e:
                    st.error(f"Could not read file: {e}")
    else:
        raw_text = st.text_area(
            "Paste document text",
            height=200,
            placeholder="Paste any business document here — reports, contracts, "
                        "meeting notes, financial summaries…",
            label_visibility="collapsed",
        )

    if raw_text and raw_text.strip():
        if st.button("Analyse Document"):
            run_analysis(raw_text, cfg)
    elif raw_text is not None and not raw_text.strip():
        st.warning("Please provide text or upload a document.")

    
    st.markdown(f"""
    <div style="margin-top:64px; padding:22px 0;
                border-top:1px solid {C['border']};
                display:flex; justify-content:space-between;
                align-items:center; flex-wrap:wrap; gap:10px">
        <div style="font-family:'Inter',sans-serif; font-size:.84rem;
                    font-weight:800; letter-spacing:-0.01em; color:{C['text_bright']}">
            AI-Powered Document Intelligence System
            <span style="color:{C['accent']}"> · </span>
            <span style="color:{C['text_muted']}; font-weight:400">Hassan Majeed</span>
        </div>
        <div style="font-size:.66rem; color:{C['text_muted']}; letter-spacing:.08em;
                    font-family:'JetBrains Mono',monospace">
            BART · KeyBERT · SpaCy · Streamlit
        </div>
    </div>
    """, unsafe_allow_html=True)


if __name__ == "__main__":
    main()