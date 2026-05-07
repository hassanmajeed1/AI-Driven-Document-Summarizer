

import pdfplumber
import PyPDF2
from docx import Document
from pathlib import Path
from loguru import logger
from typing import Optional


class DocumentExtractor:
   

    SUPPORTED_FORMATS = {'.pdf', '.docx', '.doc', '.txt'}

    def extract(self, file_path: str) -> Optional[str]:
       
        path = Path(file_path)

        
        if not path.exists():
            logger.error(f"File not found: {file_path}")
            raise FileNotFoundError(f" File nahi mili: {file_path}")

        suffix = path.suffix.lower()
        if suffix not in self.SUPPORTED_FORMATS:
            raise ValueError(
                f" Format '{suffix}' supported nahi. "
                f"Supported: {self.SUPPORTED_FORMATS}"
            )

        logger.info(f"Extracting text from: {path.name}")

       
        extractors = {
            '.pdf':  self._extract_pdf,
            '.docx': self._extract_docx,
            '.doc':  self._extract_docx,
            '.txt':  self._extract_txt,
        }

        text = extractors[suffix](str(path))
        logger.success(f"Extracted {len(text.split())} words from {path.name}")
        return text

    

    def _extract_pdf(self, file_path: str) -> str:
        
        text_parts = []

        try:
            with pdfplumber.open(file_path) as pdf:
                for page_num, page in enumerate(pdf.pages, 1):
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)
                    else:
                        logger.warning(f"Page {page_num}: No text found (possibly scanned image)")

            if text_parts:
                return "\n\n".join(text_parts)

        except Exception as e:
            logger.warning(f"pdfplumber failed: {e}. Trying PyPDF2...")

       
        text_parts = []
        with open(file_path, 'rb') as f:
            reader = PyPDF2.PdfReader(f)
            for page in reader.pages:
                text_parts.append(page.extract_text() or "")

        return "\n\n".join(text_parts)

    def _extract_docx(self, file_path: str) -> str:
       
        doc = Document(file_path)
        parts = []

        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text.strip())

        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(
                    cell.text.strip() for cell in row.cells
                    if cell.text.strip()
                )
                if row_text:
                    parts.append(row_text)

        return "\n".join(parts)

    def _extract_txt(self, file_path: str) -> str:
       
        encodings = ['utf-8', 'latin-1', 'cp1252']  
        
        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as f:
                    return f.read()
            except UnicodeDecodeError:
                continue
        
        raise ValueError(f"Could not decode file with encodings: {encodings}")